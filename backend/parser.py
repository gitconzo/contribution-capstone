import json
import re
import sys
from docx import Document

def parse_worklog_docx(file_path):
    doc = Document(file_path)
    
    # Extract hours from tables first
    hours_data = extract_hours_from_tables(doc)
    
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    weeks = []
    week_data = None
    week_header_pattern = re.compile(r"Week\s*#?\s*(\d+)", re.IGNORECASE)

    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]

        # Detect new week
        match = week_header_pattern.search(text)
        if match:
            if week_data:  # Save prior week
                weeks.append(week_data)
            week_number = int(match.group(1))
            week_data = {
                "Week": week_number,
                "TotalHours": hours_data.get(week_number),  # Get from table data
                "TasksDone": [],
                "KeyLearned": [],
                "Literature": [],
                "Issues": []
            }
            i += 1
            continue

        # Only process if week_data exists
        if week_data:
            # Total weekly hours
            if re.search(r"TOTAL WEEKLY TIME SPENT", text, re.IGNORECASE):
                # If we didn't get hours from tables, try to extract from nearby text
                if week_data["TotalHours"] is None:
                    # Look in current and next few paragraphs for hours
                    for j in range(i, min(i+3, len(paragraphs))):
                        hour_match = re.search(r"(\d+\.?\d*)", paragraphs[j])
                        if hour_match:
                            try:
                                week_data["TotalHours"] = float(hour_match.group(1))
                                break
                            except ValueError:
                                continue
                i += 1
                continue

            # Key tasks done in the week
            if re.search(r"Key tasks done|things attended", text, re.IGNORECASE):
                i += 1
                while i < len(paragraphs) and not re.search(r"Key things learned|Any literature|Issues|Plan for next week", paragraphs[i], re.IGNORECASE):
                    task_text = paragraphs[i].strip()
                    if task_text and not task_text.startswith(("o\t", "•", "o")):
                        week_data["TasksDone"].append({"Task": task_text})
                    i += 1
                continue

            # Key things learned in the week
            if re.search(r"Key things learned about Computing Technology projects", text, re.IGNORECASE):
                i += 1
                while i < len(paragraphs) and not re.search(r"Any literature|Issues|Plan for next week", paragraphs[i], re.IGNORECASE):
                    learned_text = paragraphs[i].strip()
                    if learned_text and not learned_text.startswith(("o\t", "•", "o")):
                        week_data["KeyLearned"].append(learned_text)
                    i += 1
                continue

            # Literature read in the week
            if re.search(r"Any literature read and key things learned", text, re.IGNORECASE):
                i += 1
                while i < len(paragraphs) and not re.search(r"Issues|Plan for next week", paragraphs[i], re.IGNORECASE):
                    lit_text = paragraphs[i].strip()
                    if lit_text and not lit_text.startswith(("o\t", "•", "o")):
                        # Check if this looks like literature (URLs, paper titles, etc.)
                        if (re.search(r"(http|www|\.com|\.org|\.au|arxiv|sciencedirect|youtube\.com|\.pdf)", lit_text, re.IGNORECASE) or
                            re.search(r"(paper|research|journal|proceedings|conference)", lit_text, re.IGNORECASE) or
                            len(lit_text) > 100):
                            week_data["Literature"].append(lit_text)
                        else:
                            week_data["KeyLearned"].append(lit_text)
                    i += 1
                continue

            # Issues/problems that occurred during the week
            if re.search(r"Issues/problems/Challenges|Issues and problems:|Issues/problems:", text, re.IGNORECASE):
                i += 1
                while i < len(paragraphs) and not re.search(r"Plan for next week|Summary|Week", paragraphs[i], re.IGNORECASE):
                    issue_text = paragraphs[i].strip()
                    if issue_text:
                        # Remove common bullet characters but keep the text
                        cleaned_text = re.sub(r'^[•\-\t\so]+\s*', '', issue_text)
                        if cleaned_text:
                            week_data["Issues"].append(cleaned_text)
                    i += 1
                continue

        i += 1

    # Append last week if it exists
    if week_data:
        weeks.append(week_data)

    return weeks

def extract_hours_from_tables(doc):
    """Extract hours data from tables in the document"""
    hours_data = {}
    
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text.strip() for cell in row.cells)
            
            # Look for week pattern in the entire row
            week_match = re.search(r"Week\s*#?\s*(\d+)", row_text, re.IGNORECASE)
            if week_match:
                week_num = int(week_match.group(1))
                
                # Look for hours pattern in the row cells
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    hour_match = re.search(r"^\s*(\d+\.?\d*)\s*$", cell_text)
                    if hour_match:
                        try:
                            hours = float(hour_match.group(1))
                            if 0 < hours <= 40:
                                hours_data[week_num] = hours
                                break
                        except ValueError:
                            continue
    
    # Also try to find hours in the cumulative summary table
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if "Total hours spent this week" in cell_text:
                    for next_row in table.rows[i+1:]:
                        if j < len(next_row.cells):
                            hour_text = next_row.cells[j].text.strip()
                            hour_match = re.search(r"(\d+\.?\d*)", hour_text)
                            if hour_match:
                                try:
                                    if 0 < len(next_row.cells):
                                        week_cell = next_row.cells[0].text.strip()
                                        week_match = re.search(r"(\d+)", week_cell)
                                        if week_match:
                                            week_num = int(week_match.group(1))
                                            hours = float(hour_match.group(1))
                                            hours_data[week_num] = hours
                                except (ValueError, IndexError):
                                    continue
    return hours_data

if __name__ == "__main__":
    # Get file path from command line argument
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        try:
            data = parse_worklog_docx(input_file)
            # Output JSON to stdout for Node.js to capture
            print(json.dumps(data, ensure_ascii=False))
        except Exception as e:
            print(f"Error: {str(e)}", file=sys.stderr)
            sys.exit(1)
    else:
        print("Error: No file path provided", file=sys.stderr)
        sys.exit(1)