from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import json
import re
import textstat
import spacy
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk import download as nltk_download
from difflib import get_close_matches
import argparse

nltk_download("punkt",     quiet=True)
nltk_download("punkt_tab", quiet=True)
nlp = spacy.load("en_core_web_sm")

def _count_words(sentence):
    # Count only tokens containing a letter/digit so punctuation isn't counted.
    return len([token for token in word_tokenize(sentence) if any(char.isalnum() for char in token)])

def _ends_like_sentence(sentence):
    # Punkt already detected the boundary; we just check the unit terminates
    # like a real sentence (after stripping trailing quotes/brackets/spaces).
    stripped = sentence.rstrip(" \t\n\r\"')]}>”’")
    return stripped.endswith((".", "!", "?"))

def get_text_metrics(text: str):
    cleaned = text.strip()
    if not cleaned:
        return dict(word_count=0, avg_sentence_length=0,
                    sentence_complexity=0, readability_score=0)

    sentences = sent_tokenize(cleaned)

    # Volume — prose words only (drops <3-word fragments and >60-word merged
    # table run-ons) so word_count reflects written contribution, not scaffolding.
    prose_sentences = [sentence for sentence in sentences if 3 <= _count_words(sentence) <= 60]
    word_count = sum(_count_words(sentence) for sentence in prose_sentences)

    # Quality — stricter: also require a real sentence terminator.
    quality_sentences = [sentence for sentence in prose_sentences if _ends_like_sentence(sentence)]
    quality_text = " ".join(quality_sentences)

    if quality_sentences:
        quality_word_count  = sum(_count_words(sentence) for sentence in quality_sentences)
        avg_sentence_length = quality_word_count / len(quality_sentences)
        readability         = textstat.flesch_reading_ease(quality_text)
        parsed              = nlp(quality_text)
        subordinate_deps    = {"advcl", "ccomp", "xcomp", "acl", "relcl"}
        subordinate_count   = sum(1 for token in parsed if token.dep_ in subordinate_deps)
        sentence_complexity = subordinate_count / len(quality_sentences)
    else:
        avg_sentence_length = 0.0
        readability         = 0.0
        sentence_complexity = 0.0

    return {
        "word_count":          word_count,
        "avg_sentence_length": round(avg_sentence_length, 2),
        "sentence_complexity": round(sentence_complexity, 3),
        "readability_score":   round(readability, 2),
    }


def get_heading_level(text, para):
    """
    Returns heading level 1-4, or 0 if the paragraph is not a heading.
    """
    text = text.strip()
    if not text:
        return 0

    style_name = para.style.name.lower() if para.style and para.style.name else ""

    if "heading 1" in style_name or style_name == "title":
        return 1
    if "heading 2" in style_name:
        return 2
    if "heading 3" in style_name:
        return 3
    if "heading 4" in style_name or "heading 5" in style_name or "heading 6" in style_name:
        return 4
    if "heading" in style_name:
        return 2

    if text.startswith("####"): return 4
    if text.startswith("###"):  return 3
    if text.startswith("##"):   return 2
    if text.startswith("#"):    return 1

    if re.match(r"^PART\s+\d+", text.upper()):
        return 1

    top_level = [
        "team and project plan", "acknowledgment of country",
        "team code of conduct",  "project overview",
    ]
    text_lower = text.lower()
    if any(s in text_lower for s in top_level):
        return 1

    # Numbered section headings: "1. Team Profile", "3. TEAMWORK ROADMAP"
    # ≤8 words and no commas so list items are not mistaken for headings.
    if re.match(r"^\d+[\.\)]\s+[A-Za-z]", text) and len(text.split()) <= 8 and "," not in text:
        return 2

    if "list" in style_name:
        return 0

    return 0


def _table_to_text(table):
    """Flatten a table to text: cells joined by spaces, rows by newlines."""
    row_texts = []
    for row in table.rows:
        cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cell_texts:
            row_texts.append(" ".join(cell_texts))
    return "\n".join(row_texts)


def iter_blocks(doc):
    """Yield the body's paragraphs and tables in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def extract_hierarchical_sections(doc):
    """
    Walk paragraphs AND tables in document order, identify headings, and collect
    the text between each heading (including deeper sub-sections and tables) as
    that section's content.
    Returns {title: {content, level, word_count}}.
    """
    # Ordered blocks with text + heading level (tables are never headings).
    blocks = []
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            level = get_heading_level(text, block)
            blocks.append({"text": text, "level": level, "is_heading": level > 0})
        else:  # Table
            text = _table_to_text(block)
            if text:
                blocks.append({"text": text, "level": 0, "is_heading": False})

    heading_indexes = [index for index, block in enumerate(blocks) if block["is_heading"]]

    sections = {}
    for position, heading_index in enumerate(heading_indexes):
        heading = blocks[heading_index]
        # Section ends at the next heading of equal-or-higher level; everything
        # before that (sub-headings, their bodies, and tables) belongs to it.
        section_end = len(blocks)
        for next_heading_index in heading_indexes[position + 1:]:
            if blocks[next_heading_index]["level"] <= heading["level"]:
                section_end = next_heading_index
                break

        lines   = [blocks[index]["text"]
                   for index in range(heading_index + 1, section_end)
                   if blocks[index]["text"]]
        content = "\n".join(lines).strip()
        title   = re.sub(r"^#+\s*", "", heading["text"]).strip()
        sections[title] = {
            "content":    content,
            "level":      heading["level"],
            "word_count": len(content.split()) if content else 0,
        }
    return sections


def normalize_section_name(name):
    normalized = re.sub(r"[^a-z0-9\s]+", "", name.lower())
    return re.sub(r"\s+", " ", normalized).strip()


def find_matching_section(claimed, all_sections):
    """
    Try exact → substring → fuzzy match of a claimed section name against
    the keys of all_sections.  Returns the matched key or None.
    """
    claimed_norm    = normalize_section_name(claimed)
    norm_to_orig    = {normalize_section_name(k): k for k in all_sections}

    if claimed_norm in norm_to_orig:
        return norm_to_orig[claimed_norm]

    for norm_key, orig_key in norm_to_orig.items():
        if (claimed_norm in norm_key or norm_key in claimed_norm) and len(claimed_norm) > 3:
            return orig_key

    matches = get_close_matches(claimed_norm, list(norm_to_orig), n=1, cutoff=0.5)
    if matches:
        return norm_to_orig[matches[0]]

    return None


# TABLE HELPERS

def parse_table(table):
    """Return a list of row dicts keyed by the header row."""
    headers = [c.text.strip() for c in table.rows[0].cells]
    rows    = []
    for row in table.rows[1:]:
        values = [c.text.strip() for c in row.cells]
        if any(values):
            rows.append(dict(zip(headers, values)))
    return rows


def _clean_name(name):
    """Strip parentheses, brackets, asterisks, and extra whitespace."""
    return re.sub(r"[()[\]*]", "", str(name or "")).strip()


def _get_contribution_text(row):
    """
    Return the contribution description from a row dict, trying several
    known column-name variants before falling back to any key that contains
    'contribution' or 'statement'.  Different cohort documents use different
    headings for the same field.
    """
    for key in (
        "Description of contribution in team and project planning",
        "Statement of contribution to the report",
        "Statement of Contribution to the Report",
        "Contribution",
        "Description",
    ):
        val = row.get(key, "")
        if val:
            return val.strip()

    for key, val in row.items():
        if key and ("contribution" in key.lower() or "statement" in key.lower()):
            return (val or "").strip()

    return ""


# PER-STUDENT TABLE EXTRACTION

def _extract_profile_rows(table, per_student_profile, skip_header=True):
    """
    Read a Team Profile table and merge each student's row text into
    per_student_profile.

    skip_header=False is used for continuation tables that Word creates when
    a large table spans a page break - those tables have no header row.
    The last-seen name is carried forward to handle multi-row merged cells.
    """
    rows      = table.rows[1:] if skip_header else table.rows
    last_name = None
    for row in rows:
        cells = [c.text.strip() for c in row.cells]
        if not cells:
            continue
        if cells[0]:
            last_name = _clean_name(cells[0])
        if not last_name:
            continue
        row_text = " ".join(c for c in cells[1:] if c)
        if row_text:
            if last_name in per_student_profile:
                per_student_profile[last_name] += " " + row_text
            else:
                per_student_profile[last_name] = row_text


def _extract_role_rows(table, header_cells):
    """
    Read a Team Role table and return {clean_name: "role justification text"}.
    """
    name_idx = next((i for i, h in enumerate(header_cells)
                     if "student" in h.lower() or "name" in h.lower()), 0)
    role_idx = next((i for i, h in enumerate(header_cells)
                     if "role" in h.lower() and i != name_idx), 1)
    just_idx = next((i for i, h in enumerate(header_cells)
                     if "justif" in h.lower()), -1)

    per_student_role = {}
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if not cells[name_idx]:
            continue
        name  = _clean_name(cells[name_idx])
        parts = []
        if role_idx < len(cells) and cells[role_idx]:
            parts.append(cells[role_idx])
        if 0 <= just_idx < len(cells) and cells[just_idx]:
            parts.append(cells[just_idx])
        if name and parts:
            per_student_role[name] = " ".join(parts)
    return per_student_role


# TABLE IDENTIFICATION 
_KNOWN_TABLE_KEYWORDS = {
    "technical skills", "student", "team role", "justif",
    "contribution",     "mitigation", "impact", "story", "functional",
}


def _extract_tables(doc):
    """
    First pass: classify every table in the document and extract per-student
    profile and role content.

    Returns:
        tables              – dict matching result["tables"] schema
        per_student_profile – {clean_name: profile_text}
        per_student_role    – {clean_name: role_text}
    """
    tables              = {}
    per_student_profile = {}
    per_student_role    = {}
    team_profile_ncols  = 0
    claimed_ids         = set()

    def _append(key, include_in_metrics, table_obj):
        if key not in tables:
            tables[key] = {"include_in_metrics": include_in_metrics, "rows": []}
        tables[key]["rows"].extend(parse_table(table_obj))

    for idx, table in enumerate(doc.tables):
        if not table.rows:
            continue
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        header       = " ".join(h.lower() for h in header_cells)

        if "student" in header and "contribution" in header:
            _append("TeamContributions", True, table)
            claimed_ids.add(idx)

        elif "technical skills" in header:
            _append("TeamProfile", False, table)
            team_profile_ncols = len(header_cells)
            claimed_ids.add(idx)
            _extract_profile_rows(table, per_student_profile, skip_header=True)

        elif "team role" in header and ("justif" in header or "student" in header):
            _append("TeamRole", False, table)
            claimed_ids.add(idx)
            per_student_role.update(_extract_role_rows(table, header_cells))

        elif "impact on project" in header or "mitigation" in header:
            _append("RiskMitigation", True, table)
            claimed_ids.add(idx)

        elif "user story" in header and "priority" in header:
            _append("ProductBacklog", True, table)
            claimed_ids.add(idx)

        elif "functional requirements" in header or "story" in header:
            _append("HighLevelRequirements", True, table)
            claimed_ids.add(idx)

    # Second pass: Word splits large tables across page breaks into separate
    # table objects.  Pick up unclaimed tables whose column count matches the
    # Team Profile table - these are continuation pages with no header row.
    if team_profile_ncols > 0:
        for idx, table in enumerate(doc.tables):
            if idx in claimed_ids or not table.rows:
                continue
            h_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            if len(h_cells) != team_profile_ncols:
                continue
            if any(kw in " ".join(h_cells) for kw in _KNOWN_TABLE_KEYWORDS):
                continue
            _extract_profile_rows(table, per_student_profile, skip_header=False)

    return tables, per_student_profile, per_student_role


# STUDENT ATTRIBUTION

def _build_student_records(contributions, per_student_profile, per_student_role, sections):
    """
    For each student row in the contributions table, match their claimed
    sections/tables to extracted text and compute metrics.
    Returns {student_name: {sections_written, raw_text, metrics}}.
    """
    students = {}

    for row in contributions:
        name = (row.get("Student Name") or row.get("Student name") or "").strip()
        if not name:
            continue
        desc = _get_contribution_text(row)
        if not desc:
            continue

        # Split on newlines; if still a single item, split on commas/semicolons
        claimed_parts = [s.strip() for s in re.split(r"[\n\r]+", desc) if s.strip()]
        if len(claimed_parts) == 1:
            claimed_parts = [s.strip() for s in re.split(r"[,;]+", claimed_parts[0]) if s.strip()]

        clean_name         = _clean_name(name)
        section_text_parts = []
        matched_sections   = []

        for claim in claimed_parts:
            claim_lower = claim.lower()
            handled     = False

            # Team Profile: use per-student table row rather than the section
            # heading (which is just a table caption with no useful prose).
            if re.search(r"team.{0,5}profile|own.{0,5}profile", claim_lower):
                profile_text = per_student_profile.get(name) or per_student_profile.get(clean_name)
                if profile_text and "Team Profile (personal row)" not in matched_sections:
                    section_text_parts.append(profile_text)
                    matched_sections.append("Team Profile (personal row)")
                handled = True  # always - the section heading has no useful text

            # Team Role: checked independently so a combined claim like
            # "own team profile and role justification" matches both patterns.
            if re.search(r"role.{0,10}justif|team.{0,5}role|role.{0,10}desc", claim_lower):
                role_text = per_student_role.get(name) or per_student_role.get(clean_name)
                if role_text and "Team Role (personal row)" not in matched_sections:
                    section_text_parts.append(role_text)
                    matched_sections.append("Team Role (personal row)")
                    handled = True

            # General prose sections (Roadmap, Document Management, Risk…)
            if not handled:
                matched = find_matching_section(claim, sections)
                if matched:
                    matched_sections.append(matched)
                    section_text_parts.append(sections[matched])

        combined_text    = "\n\n".join(section_text_parts).strip()
        students[name]   = {
            "sections_written": matched_sections,
            "raw_text":         combined_text,
            "metrics":          get_text_metrics(combined_text),
        }

    return students


# MAIN PARSER
def parse_project_plan_docx(docx_path, output_json_path=None):
    doc = Document(docx_path)

    tables, per_student_profile, per_student_role = _extract_tables(doc)

    sections_data = extract_hierarchical_sections(doc)
    sections      = {k: v["content"] for k, v in sections_data.items()}

    contributions = tables.get("TeamContributions", {}).get("rows", [])
    students      = _build_student_records(
        contributions, per_student_profile, per_student_role, sections
    )

    result = {
        "source_file":     Path(docx_path).name,
        "tables":          tables,
        "students":        students,
        "overall_metrics": get_text_metrics("\n\n".join(sections.values())),
    }

    if output_json_path:
        with open(output_json_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

    return result


# CLI

def _filter_students_to_roster(res_obj, roster_names):
    keep = {
        name: details
        for name, details in res_obj.get("students", {}).items()
        if name in roster_names or get_close_matches(name, roster_names, n=1, cutoff=0.85)
    }
    res_obj["students"] = keep
    return res_obj


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("input_path")
    ap.add_argument("output_path", nargs="?")
    ap.add_argument("--students-json", default=None)
    args = ap.parse_args()

    out_path = args.output_path or str(Path(args.input_path).with_suffix(".json"))
    res      = parse_project_plan_docx(args.input_path, out_path)

    print(f"\nSections found: {list(res.get('tables', {}).keys())}")
    print(f"Students attributed: {list(res.get('students', {}).keys())}")
    for name, data in res["students"].items():
        print(f"  {name}: {data['metrics']['word_count']} words, "
              f"sections={data['sections_written']}")

    if args.students_json:
        try:
            with open(args.students_json, "r", encoding="utf-8") as f:
                roster = [(s.get("name") or "").strip() for s in json.load(f)
                          if (s.get("name") or "").strip()]
        except Exception:
            roster = None
        if roster:
            res = _filter_students_to_roster(res, roster)
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(res, f, indent=2, ensure_ascii=False)

    print(f"\nProject Plan parsed --> {out_path}")
