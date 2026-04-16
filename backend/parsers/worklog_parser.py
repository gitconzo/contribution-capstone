import json
import re
import sys
from docx import Document


def parse_worklog_pdf(file_path):
    from pdfminer.high_level import extract_text

    raw_text = extract_text(file_path)
    lines = raw_text.splitlines() if raw_text else []
    paragraphs = [line.strip() for line in lines if line.strip()]
    return _parse_paragraphs(paragraphs, hours_data={})


def parse_worklog_docx(file_path):
    doc = Document(file_path)
    
    # Extract hours from tables first
    hours_data = extract_hours_from_tables(doc)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return _parse_paragraphs(paragraphs, hours_data)


def _is_number(s):
    try:
        float(s)
        return True
    except (ValueError, TypeError):
        return False


def _parse_paragraphs(paragraphs, hours_data):
    weeks = []
    week_data = None
    week_header_pattern = re.compile(r"Week\s*#?\s*(\d+)", re.IGNORECASE)
    hours_inline_pattern = re.compile(r"(?:total\s*hours?|hours?\s*worked)[^\d]*(\d+\.?\d*)", re.IGNORECASE)
    total_weekly_pattern = re.compile(r"TOTAL\s+WEEKLY\s+TIME\s+SPENT", re.IGNORECASE)

    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]

        # Detect new week
        match = week_header_pattern.search(text)
        if match:
            if week_data:  # Save prior week
                weeks.append(week_data)
            week_number = int(match.group(1))
            week_data = {
                "Week": week_number,
                "TotalHours": hours_data.get(week_number),
                "TasksDone": [],
                "KeyLearned": [],
                "Literature": [],
                "Issues": []
            }
            i += 1
            continue

        if week_data:
            # PDF format: "TOTAL WEEKLY TIME SPENT" followed by numbers on subsequent lines.
            # The last consecutive number is the weekly total (preceding numbers are sub-totals).
            if total_weekly_pattern.search(text):
                i += 1
                last_num = None
                while i < len(paragraphs) and _is_number(paragraphs[i]):
                    last_num = float(paragraphs[i])
                    i += 1
                if last_num is not None and week_data["TotalHours"] is None:
                    week_data["TotalHours"] = last_num
                continue

            # Inline hours line (common in PDF-exported worklogs)
            hours_match = hours_inline_pattern.search(text)
            if hours_match and week_data["TotalHours"] is None:
                try:
                    week_data["TotalHours"] = float(hours_match.group(1))
                except ValueError:
                    pass
                i += 1
                continue

            # Key sections detection
            if re.search(r"Key tasks done|things attended", text, re.IGNORECASE):
                i += 1
                while i < len(paragraphs) and not re.search(r"Key things learned|Any literature|Issues|Plan for next week", paragraphs[i], re.IGNORECASE):
                    week_data["TasksDone"].append(paragraphs[i])
                    i += 1
                continue

            if re.search(r"Key things learned", text, re.IGNORECASE):
                i += 1
                while i < len(paragraphs) and not re.search(r"Any literature|Issues|Plan for next week", paragraphs[i], re.IGNORECASE):
                    week_data["KeyLearned"].append(paragraphs[i])
                    i += 1
                continue

            if re.search(r"Any literature", text, re.IGNORECASE):
                i += 1
                while i < len(paragraphs) and not re.search(r"Issues|Plan for next week", paragraphs[i], re.IGNORECASE):
                    week_data["Literature"].append(paragraphs[i])
                    i += 1
                continue

            if re.search(r"Issues|Challenges", text, re.IGNORECASE):
                i += 1
                while i < len(paragraphs) and not re.search(r"Plan for next week|Week", paragraphs[i], re.IGNORECASE):
                    week_data["Issues"].append(paragraphs[i])
                    i += 1
                continue

        i += 1

    if week_data:
        weeks.append(week_data)

    return weeks

def extract_hours_from_tables(doc):
    hours_data = {}
    for table in doc.tables:
        for row in table.rows:
            match = re.search(r"Week\s*#?\s*(\d+)", row.cells[0].text, re.IGNORECASE)
            if match:
                try:
                    week = int(match.group(1))
                    hours = float(re.findall(r"\d+\.?\d*", row.cells[-1].text)[0])
                    hours_data[week] = hours
                except:
                    pass
    return hours_data

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 worklog_parser.py input.docx output.json", file=sys.stderr)
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    try:
        if input_file.lower().endswith(".pdf"):
            data = parse_worklog_pdf(input_file)
        else:
            data = parse_worklog_docx(input_file)
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"Saved JSON to {output_file}")
    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)
