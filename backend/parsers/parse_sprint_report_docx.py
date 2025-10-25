import json
import re
import sys
from pathlib import Path
from docx import Document

def extract_text_from_docx(doc_path):
    # Extract paragraphs from the Word document as plain text
    doc = Document(doc_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return lines

def extract_section(lines, start_pattern, end_patterns):
    # Extract lines between start and next matching end pattern
    text = "\n".join(lines)
    start_match = re.search(start_pattern, text, re.IGNORECASE)
    if not start_match:
        return ""
    start_idx = start_match.end()

    end_idx = len(text)
    for ep in end_patterns:
        match = re.search(ep, text[start_idx:], re.IGNORECASE)
        if match:
            end_idx = start_idx + match.start()
            break

    return text[start_idx:end_idx].strip()

def parse_attendance_table(doc):
    # Attendance table data extraction
    tables = doc.tables
    attendance_data = []

    for table in tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if any("Contribution" in h for h in headers) and any("Respect" in h for h in headers):
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 9 and cells[0]:
                    attendance_data.append({
                        "Student": cells[0],
                        "Contribution": cells[1],
                        "Initiatives": cells[2],
                        "Communication": cells[3],
                        "Respect": cells[8]
                    })
    return attendance_data

def parse_contributions(text):
    # Extraction of indiv student contribution sections
    pattern = r"(Connor Lack|Jason Vo|Jen Mao|Kavindu Bhanuka Weragoda|Md Hridoy Mia)"
    matches = list(re.finditer(pattern, text))
    contributions = {}

    for i, m in enumerate(matches):
        name = m.group(1)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start:end].strip()

        # Ignore content that could look like sprint section headings
        if not content or re.match(r"^\s*(\d+\.\s*Sprint|Sprint\s+plan)", content, re.IGNORECASE):
            contributions[name] = ""
        else:
            contributions[name] = content

    return contributions

def parse_sprint_report_docx(docx_path):
    # main parser
    doc = Document(docx_path)
    lines = extract_text_from_docx(docx_path)
    text = "\n".join(lines)

    attendance = parse_attendance_table(doc)
    contributions = parse_contributions(text)
    sprint_plan = extract_section(lines, r"1\.\s*Sprint plan", [r"2\.\s*Sprint progress"])
    sprint_progress = extract_section(lines, r"2\.\s*Sprint progress", [r"3\.\s*Sprint Demonstration"])
    sprint_demo = extract_section(lines, r"3\.\s*Sprint Demonstration", [r"4\.\s*Sprint review"])
    sprint_review = extract_section(lines, r"4\.\s*Sprint review", [r"5\.\s*Retrospect"])
    retrospect = extract_section(lines, r"5\.\s*Retrospect", [r"6\.\s*Lessons learned"])
    lessons = extract_section(lines, r"6\.\s*Lessons learned", [r"\Z"])

    result = {
        "attendance_table": attendance,
        "contributions": contributions,
        "sprint_plan": sprint_plan,
        "sprint_progress": sprint_progress,
        "sprint_demonstration": sprint_demo,
        "sprint_review": sprint_review,
        "retrospect": retrospect,
        "lessons_learned": lessons,
    }

    out_path = Path(docx_path).with_name("sprint_report_summary.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    print(f"Parsed data saved to {out_path}")
    return result


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python parse_sprint_report_docx.py <report.docx>")
        sys.exit(1)
    file_path = sys.argv[1]
    parse_sprint_report_docx(file_path)
