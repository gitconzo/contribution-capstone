import sys
import json
import re
from docx import Document

 
def clean_name(name):
    return (name or "").strip().lower()
 
def is_number(val):
    try:
        float(val)
        return True
    except (ValueError, TypeError):
        return False

def names_match(a, b):
    # Check if names share the same first word (first name)
    a_first = a.split()[0] if a.split() else a
    b_first = b.split()[0] if b.split() else b
    return a_first == b_first or a == b
 
def parse_peer_review(docx_path):
    doc = Document(docx_path)
    results = {}
    reviewer_name = None
 
    for para in doc.paragraphs:
        text = para.text.strip()
        if "your name" in text.lower():
            match = re.search(r'your name[:\s]+(.+)', text, re.IGNORECASE)
            if match:
                reviewer_name = match.group(1).strip()
                break
 
    if not reviewer_name:
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                for i, cell in enumerate(cells):
                    if "your name" in cell.lower() and i + 1 < len(cells):
                        reviewer_name = cells[i + 1].strip()
                        break
 
    # find the scoring table
    scoring_table = None
    for table in doc.tables:
        headers = [c.text.strip().upper() for c in table.rows[0].cells]
        # check if headers contain A through J
        if all(letter in headers for letter in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J']):
            scoring_table = table
            break
 
    if not scoring_table:
        print(f"Warning: Could not find scoring table in {docx_path}")
        return {"reviewer": reviewer_name, "scores": {}}
 
    # Get column indices for A-J
    headers = [c.text.strip().upper() for c in scoring_table.rows[0].cells]
    score_cols = {}
    for letter in 'ABCDEFGHIJ':
        if letter in headers:
            score_cols[letter] = headers.index(letter)
 
    name_col = None
    for i, h in enumerate(headers):
        if h in ("NAME", "TEAM MEMBER NAMES", "TEAM MEMBER NAMES (INCLUDING YOURSELF)"):
            name_col = i
            break
    if name_col is None:
        name_col = 1
 
    reviewer_clean = clean_name(reviewer_name)
 
    for row in scoring_table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if not cells or not cells[name_col]:
            continue
 
        student_name = cells[name_col].strip()
        if not student_name or student_name == "0":
            continue
 
        student_clean = clean_name(student_name)
 
        # skip self-assessment
        if reviewer_clean and student_clean and names_match(reviewer_clean, student_clean):
            continue
 
        # sum A-J scores
        total = 0
        valid = False
        for letter, col_idx in score_cols.items():
            if col_idx < len(cells) and is_number(cells[col_idx]):
                val = float(cells[col_idx])
                if val > 0:
                    valid = True
                total += val
 
        if not valid:
            continue
 
        if student_clean not in results:
            results[student_clean] = []
        results[student_clean].append(total)
 
    return {
        "reviewer": reviewer_name,
        "scores": results
    }
 
def main():
    if len(sys.argv) < 3:
        print("Usage: python parse_peer_review.py <input.docx> <output.json>")
        sys.exit(1)
 
    docx_path = sys.argv[1]
    output_path = sys.argv[2]
 
    result = parse_peer_review(docx_path)
 
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2)
 
    print(f"Peer review parsed: {len(result['scores'])} students scored")
 
if __name__ == "__main__":
    main()